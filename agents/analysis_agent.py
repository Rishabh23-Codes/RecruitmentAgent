import re
import PyPDF2
import io
from langchain_groq import ChatGroq
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from concurrent.futures import ThreadPoolExecutor
import tempfile
import os
import json
import time
from config import GROQ_API_KEY,LLM_MODEL
import docx
from docling.document_converter import DocumentConverter
# from langchain_ollama.chat_models import ChatOllama
from ui_utils import role_requirements as require
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor


def safe_llm_invoke(llm, prompt, max_retries=3):
    for attempt in range(max_retries):
        try:
            return llm.invoke(prompt)
        except Exception as e:
            if "rate limit" in str(e).lower() or "429" in str(e):
                sleep_time = 1 + attempt * 1.5
                time.sleep(sleep_time)
                continue
            raise e
    raise RuntimeError("Groq API failed after multiple retries")


class ResumeAnalysisAgent:
    def __init__(self):
        self.api_key=GROQ_API_KEY
        self.cutoff_score=75
        self.resume_text=None
        self.rag_vectorstore=None
        self.analysis_result=None
        self.jd_text=None
        self.extracted_skills=None
        self.resume_weaknesses=[]
        self.resume_strengths=[]
        self.improvement_suggestions={}
        self.skills=[]
        self.education=[]
        self.experience=[]
        self.job_id=None
        self.contact_info={"email":"","phone":""}
        self.llm = ChatGroq(model=LLM_MODEL, api_key=GROQ_API_KEY)
        # self.llm=ChatOllama(model=LLM_MODEL,temperature=0)
        self.embeddings=HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')


    # def extract_text_from_pdf(self,pdf_file):
    #     """Extract text from a PDF file"""
    #     try:
    #         if hasattr(pdf_file,'getvalue'):
    #             pdf_data=pdf_file.getvalue()
    #             pdf_file_like=io.BytesIO(pdf_data)
    #             reader=PyPDF2.PdfReader(pdf_file_like)
    #         else:
    #             reader=PyPDF2.PdfReader(pdf_file)
    #         text=""
    #         for page in reader.pages:
    #             text+=page.extract_text()

    #         return text
    #     except Exception as e:
    #         print(f"Error extracting text from PDF: {e}")
    #         return ""
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from a PDF file using Docling"""
        try:
            converter = DocumentConverter()

            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_file:
                # Write the PDF content to the temp file
                if hasattr(pdf_file, "getvalue"):
                    temp_file.write(pdf_file.getvalue())
                else:
                    # In case pdf_file is already bytes
                    temp_file.write(pdf_file)
                temp_path = temp_file.name
            
            # Convert the PDF to text
            result = converter.convert(temp_path).document
            response=result.export_to_text()
            return response

        except Exception as e:
            print(f"Error extracting text from PDF with Docling: {e}")
            return ""
        
    def extract_text_from_docx(self,docx_file):
        """Extract text from a docx file"""
        try:
            if hasattr(docx_file, "getvalue"):
                file_stream = io.BytesIO(docx_file.getvalue())
                doc = docx.Document(file_stream)
            else:
                doc = docx.Document(docx_file)
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
            return extracted_text
        except Exception as e:
            print(f"Error extracting text from docx file: {e}")
            return ""

        

    def extract_text_from_txt(self,txt_file):
        """Extract text from a text file"""
        try:
            if hasattr(txt_file,'getvalue'):
                return txt_file.getvalue().decode('utf-8')
            else:
                with open(txt_file,'r',encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            print(f"Error extracting text from text file: {e}")
            return ""
        

        
    def extract_text_from_file(self,file):
        """Extract text from a file (PDF or TXT)"""
        if hasattr(file,'name'):
            file_extension=file.name.split('.')[-1].lower()
        else:
            file_extension=file.split('.')[-1].lower()

        if file_extension=='pdf':
            return self.extract_text_from_pdf(file)
        elif file_extension=='txt':
            return self.extract_text_from_txt(file)
        elif file_extension=='docx':
            return self.extract_text_from_docx(file)
        else:
            print(f"Unsupported file extension: {file_extension}")
            return "" 
        
    def create_rag_vector_store(self,text):
        """Create a vector store for RAG"""
        text_splitter=RecursiveCharacterTextSplitter(
            chunk_size=1600,
            chunk_overlap=200,
            length_function=len,
        )
        chunks=text_splitter.split_text(text)
        vectorstore=FAISS.from_texts(chunks,self.embeddings)
        return vectorstore
    
    def create_vector_store(self,text):
        """Create a simpler vector store for skill analysis"""
        vectorstore=FAISS.from_texts([text],self.embeddings)
        return vectorstore
    


####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------
####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------




    def extract_info_from_resume(self,resume_text):
        try:
            prompt = f"""
            System: You are a resume parsing expert.
            Task: Extract all valid information of the following fields (skills, education, experience) from the variable "Resume Content" Only. Only extract information present in "Resume Content". Do NOT hallucinate. Each field's value must be a list of strings. If no information is found, return ["Not found"].
            Note: Use {resume_text} wherever the "Resume Content" needs to be referenced.

            Fields to extract:

            1. Skills: Technical tools, programming languages, frameworks, and domain knowledge. Example: ["Python", "JavaScript", "React.js", "Node.js", "SQL", "Docker", "LangChain", "Machine Learning", "llama3:1b", "RAG", "AI-Agent"]
            2. Education: Degree, major/subject, institution, and explicitlty graduation year . Example: ["B.Tech in Computer Science, Stanford University (2020–2024)", "B.Sc in Data Science, University of California, Berkeley (2021–2024)", "B.E in IT, IIT Bombay (2019–2023)"]
            3. Experience: Hands-on experience from projects, internships, hackathons, open-source contributions, competitions, or work experience. Example: ["Built an AI agent using LangChain and ChatGroq (Llama-3.1-8B) for real-time query handling and structured responses.", "Developed a recommendation system with Python and Scikit-learn to provide personalized product suggestions.", "Created an NLP pipeline with SpaCy and Transformers for sentiment analysis and entity recognition on large datasets.", "Implemented a serverless REST API using FastAPI and AWS Lambda for high-volume requests with low latency.", "Designed an interactive dashboard with Streamlit and Plotly to visualize real-time sales data and key metrics."]

            Expected Output example:
            {{
                "skills": ["Python", "React.js", "Docker"],
                "education": ["B.Tech in Computer Science, Stanford University (2020–2024)"],
                "experience": ["Hackathon Experience (college or national-level events)", "Open-Source Contribution Experience (GitHub projects)"]
            }}

                **RULES:**
                - Only use information explicitly mentioned in the Resume
                - If no information found in any field: ["None found"]
                - Return ONLY valid JSON - no explanations!

            """
            response_text = safe_llm_invoke(self.llm, prompt).content.strip()
            print(f"📄 Resume LLM response preview: {response_text[:10]}...") 
            try:
                llm_data = json.loads(response_text)
            except json.JSONDecodeError:
                # Fallback: extract JSON block safely
                json_start = response_text.find("{")
                json_end = response_text.rfind("}") + 1
                if json_start != -1 and json_end != -1:
                    llm_data = json.loads(response_text[json_start:json_end])
                else:
                    print("❌ LLM response does not contain JSON")
                    llm_data = {}

            skills = llm_data.get("skills", ["Not found"])
            education = llm_data.get("education", ["Not found"])
            experience = llm_data.get("experience", ["Not found"])
            print(f"✅ Extracted: {len(skills)} skills, {len(education)} education, {len(experience)} experience")
            return skills, education, experience

        except Exception as e:
            print(f"Error extracting skills/education/experience : {e}")

         # Fallback
        return ["Not found"], ["Not found"], ["Not found"]
    




        

####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------
####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------
####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------

    def compare_resume_jd_new(self,skills,experience,education,role_requirements=None,custom_jd=None):
            try:
                context_text=""
                experiences=",".join(experience)[:2000]
                skill=",".join(skills)
                role=None
                if custom_jd:
                    jd_text=self.extract_text_from_file(custom_jd)
                    jd_vectorstore=self.create_rag_vector_store(jd_text)
                    retriever=jd_vectorstore.as_retriever(search_kwargs={"k": 3})
                    query="Extract all technical skills, programming languages, frameworks, tools, cloud platforms, databases, and relevant technologies mentioned in this job description. Include both mandatory and optional skills.For example: ['Python', 'JavaScript', 'React.js', 'Node.js', 'SQL', 'Docker', 'AWS', 'Machine Learning', 'LangChain']."
                    relevant_chunks = retriever.invoke(query)
                    context_text = "\n".join([doc.page_content for doc in relevant_chunks])[:3000]
                    print(f"✅ JD context: {len(relevant_chunks)} chunks")
                elif role_requirements:
                    role = next((key for key, values in require.items() if any(v in values for v in role_requirements)), None)
                    context_text=",".join(role_requirements)
                    print(f"✅ Role requirements: {len(role_requirements)} skills")

                if not context_text:
                    print("❌ NO JD CONTEXT - Returning empty!")

                # Prompt to extract technical skills in strict JSON format
                prompt = f"""
            
                System: You are an expert resume analysis and job-alignment specialist.

                The resume information is already extracted and provided as structured inputs
                (skills, experience, education). Use these inputs directly without re-interpreting
                or re-extracting resume data.

                Your task is to compare the provided resume details with the job context,analysis
                identify alignment, gaps, weaknesses, and generate actionable improvement insights.

                All analysis must be strictly grounded in the given inputs.
                Do not assume or infer information not explicitly present.

                ====================================================
                INPUT DATA
                ====================================================

                === RESUME DETAILS ===
                Skills:
                {skill}

                Education:
                {education}

                Experience:
                {experiences}

                === JOB CONTEXT ===
                Job detail in job_context:
                {context_text}

                ====================================================
                ANALYSIS OBJECTIVES
                A) JD–RESUME COMPARISON
                ====================================================
                

                1. Identify skills that align between the resume and job_context.
                2. Explain how the resume demonstrates the aligned skills.
                3. Identify missing skills required by the job.
                4. Extract all skills mentioned in the job context.
                5. Determine the job role and a concise job description.
                6. Provide resume-focused assessment, improvement suggestions, and ATS guidance.
                7. Analyze weaknesses ONLY for missing skills and provide improvement actions.

                ====================================================
                FIELDS TO RETURN (KEYS MUST MATCH EXACTLY)
                ====================================================

                1. Matching Skills  
                - Skills present in resume skills or experience
                - AND explicitly mentioned in job_context

                2. Skill Reasoning  
                - One-to-one correspondence with Matching Skills (same order)
                - Concise explanation of evidence from experience or skills

                3. Missing Skills  
                - Skills explicitly required or mentioned in job_context
                - Not present in resume skills or experience
                - Ordered them in decreasing priority based on their importance in the job description

                4. Extracted Skills  
                - All skills explicitly mentioned in job_context

                5. Job Role  
                - {role if role else "extract the job role from job_context"}

                6. Job Description  
                - {"Generate a short and concise professional Job description based on this role: "+role if role else "Extract the short and concise job description from the 'job_context' only"}

                Example:  
                    Resume Skills: "Python, React.js, SQL, Docker"  
                    Job Description Context: "Looking for candidates with Python, JavaScript, React.js, Node.js, SQL, AWS, Docker, and CI/CD experience."  

                    Expected example structure:
                    {{
                        "Matching Skills": ["Python", "React.js", "SQL", "Docker"],
                        "Skill Reasoning": [
                            "Python: Strong experience indicated in multiple projects",
                            "React.js: Used in personal web development projects",
                            "SQL: Demonstrated through database management tasks",
                            "Docker: Used for containerizing projects"
                        ],
                        "Missing Skills": ["JavaScript", "Node.js", "AWS", "CI/CD"],
                        "Extracted Skills": ["Python", "JavaScript", "React.js", "Node.js", "SQL", "AWS", "Docker", "CI/CD"],
                        "Job Role": "Web Developer",
                        "Job Description": "Looking for candidates with Python, JavaScript, React.js, Node.js, SQL, AWS, Docker, and CI/CD experience."  
                    }}
                ====================================================
                B) RESUME OVERALL ANALYSIS (Resume-Focused)
                ====================================================

                Using ONLY resume data and comparison results, generate:

                1. Overall Assessment  
                - Strengths (derived from matching skills)
                - Areas for improvement
                - Suitable job roles or industry sectors

                2. Content Improvements  
                - Ways to quantify achievements
                - Skills presentation improvement
                - Missing critical skill

                3. Format Suggestions  
                - Structure improvement
                - Length suggestion
                - Readability tip

                4. ATS Optimization  
                - Additional keywords
                - Formatting pitfall
                - File format recommendation

                Expected structure example:

                "Resume Overall Analysis": {{
                    "overall_assessment": {{
                        "strengths": ["", "", ""],
                        "areas_for_improvement": ["", ""],
                        "suitable_roles_or_sectors": ["", ""]
                    }},
                    "content_improvements": {{
                        "quantification_suggestions": ["", ""],
                        "skills_presentation": [""],
                        "missing_critical_skill": [""]
                    }},
                    "format_suggestions": {{
                        "structure": [""],
                        "length": [""],
                        "readability": [""]
                    }},
                    "ats_optimization": {{
                        "additional_keywords": ["", "", ""],
                        "formatting_pitfall": [""],
                        "file_format_recommendation": [""]
                    }}
                }}

                Guidelines:
                - Concise, actionable points
                - Maximum 3 bullets per subsection
                - Plain bullet text only
                - No special characters or symbols

                ====================================================
                C) DETAILED WEAKNESS & IMPROVEMENT ANALYSIS
                ====================================================

                Analyze weaknesses ONLY for the Missing Skills.

                For detailed weakness and improvement analysis
                Take 5 top necessary * Missing Skills * and Do iteratively for each and every -> explain why the resume does not sufficiently demonstrate
                that skill and provide specific, realistic improvement suggestions.
                For your analysis consider:
                    1. What's missing from the resume regarding this skill?
                    2. How could it be improved with specific examples?
                    3. What specific action items would make this skill stand out?

                Return in the following example structure:

                "Detailed Weaknesses": [
                    {{
                        "skill": "<missing skill name>",
                        "weakness": "A concise description of what's missing or problematic(1–2 sentences)",
                        "improvement_suggestions": [
                            "Specific suggestion 1",
                            "Specific suggestion 2",
                            "Specific suggestion 3"
                        ]
                    }}
                ]

                ====================================================
                EXPECTED OUTPUT STRUCTURE (FINAL JSON)
                ====================================================
                The response MUST be a single valid JSON object that FILLS the following structure.
                All keys must be present exactly as shown. Replace placeholder values with required filled content.
                {{
                    "Matching Skills": [],
                    "Skill Reasoning": [],
                    "Missing Skills": [],
                    "Extracted Skills": [],
                    "Job Role": "",
                    "Job Description": "",

                    "Resume Overall Analysis": {{
                        "overall_assessment": {{
                            "strengths": [],
                            "areas_for_improvement": [],
                            "suitable_roles_or_sectors": []
                        }},
                        "content_improvements": {{
                            "quantification_suggestions": [],
                            "skills_presentation": [],
                            "missing_critical_skill": []
                        }},
                        "format_suggestions": {{
                            "structure": [],
                            "length": [],
                            "readability": []
                        }},
                        "ats_optimization": {{
                            "additional_keywords": [],
                            "formatting_pitfall": [],
                            "file_format_recommendation": []
                        }}
                    }},

                    "Detailed Weaknesses": [
                        {{
                            "skill": "",
                            "weakness": "",
                            "improvement_suggestions": []
                        }}
                    ]
                }}
                ====================================================
                STRICT RULES
                ====================================================
                - Use only information present in the inputs
                - Do not hallucinate skills, experience, or achievements
                - Maintain exact key names and nesting
                - Skill Reasoning must map 1:1 with Matching Skills
                - If no data exists for a field, return ["Not found"]
                - Return ONLY valid JSON
                - No markdown, no commentary, no explanations
                """

                response_text = safe_llm_invoke(self.llm, prompt).content.strip()
            
                # ---------- SAFE LLM JSON PARSING ----------
                print(f"📄 JD Match LLM response preview:\n{response_text[:10]}\n")

                try:
                    # 1️⃣ Try direct JSON parse
                    llm_data = json.loads(response_text)

                except json.JSONDecodeError:
                    # 2️⃣ Fallback: extract JSON block manually
                    json_start = response_text.find("{")
                    json_end = response_text.rfind("}") + 1

                    if json_start != -1 and json_end != -1:
                        try:
                            llm_data = json.loads(response_text[json_start:json_end])
                        except json.JSONDecodeError as e:
                            print("❌ JD Match JSON extraction failed:", e)
                            llm_data = {}
                    else:
                        print("❌ No JSON found in JD Match LLM response")
                        llm_data = {}

                # ---------- SAFE FIELD EXTRACTION ----------
                matching_skills = llm_data.get("Matching Skills", [])
                skill_reasoning = llm_data.get("Skill Reasoning", [])
                missing_skills = llm_data.get("Missing Skills", [])
                extracted_skills = llm_data.get("Extracted Skills", [])
                job_role=llm_data.get("Job Role","")
                job_description=llm_data.get("Job Description","")
                resume_overall_analysis=llm_data.get("Resume Overall Analysis",{})
                detailed_weaknesses=llm_data.get("Detailed Weaknesses",[])

                print(
                    f"✅ LLM Parsed → "
                    f"Matches: {len(matching_skills)}, "
                    f"JD Skills: {len(extracted_skills)}, "
                    f"Gaps: {len(missing_skills)}, "
                    f"Role: {job_role}, "
                    f"JD Desc Len: {len(job_description)}, "
                    f"Analysis Sections: {len(resume_overall_analysis)}, "
                    f"Weaknesses: {len(detailed_weaknesses)}"
                )

                
                # ✅ RULE-BASED CALCULATIONS (No semantic analysis needed)
                total_jd_skills = len(extracted_skills)
                match_count = len(matching_skills)
                overall_score = int((match_count / max(1, total_jd_skills)) * 100)
                
                # Strengths = matching_skills (as requested)
                strengths = matching_skills.copy()
                improvement_area = missing_skills if overall_score < self.cutoff_score else []
                
                # PERFECT OUTPUT STRUCTURE (Exactly as requested)
                result = {
                    "resume_skills": skills,   
                    "experience": experience,
                    "education": education,           # Input parameter
                    "matching_skills": matching_skills,          # LLM + strengths
                    "strengths": strengths,                      # = matching_skills
                    "skill_reasoning": skill_reasoning,          # LLM
                    "missing_skills": missing_skills,            # LLM
                    "extracted_skills": extracted_skills,        # LLM (JD skills)
                    "overall_score": overall_score,              # = match_percentage
                    "improvement_area":improvement_area,
                    "selected": overall_score >= self.cutoff_score,
                    "job_role": role if role else job_role,
                    "job_description":job_description,
                    "resume_overall_analysis":resume_overall_analysis,
                    "detailed_weaknesses":detailed_weaknesses
                }
                
                print(f"✅ Analysis: {overall_score}% | Strengths: {len(strengths)} | Gaps: {len(missing_skills)}")
                return result
                
            except Exception as e:
                print(f"❌ compare_resume_jd error: {e}")
                return {
                    "resume_skills": skills,
                    "experience": experience,
                    "education": education,
                    "matching_skills": [],
                    "strengths": [],
                    "skill_reasoning": [],
                    "missing_skills": [],
                    "extracted_skills": [],
                    "overall_score": 0,
                    "improvement_area": [],
                    "selected": False,
                    "job_role": "",
                    "job_description":"",
                    "resume_overall_analysis":{},
                    "detailed_weaknesses":[],
                    "error": str(e)
                }









####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------
####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------
####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------




####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------
####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------
####-------------------------------------------####-------------------------------------------####-------------------------------------------####-------------------------------------------

    def preprocess_resume(self, resume_file, job_id):
            """
            Runs in background. Stores everything on self.
            """
            # 🔒 Set active job
            self.job_id = job_id

            self.resume_text = self.extract_text_from_file(resume_file)

            with ThreadPoolExecutor(max_workers=3) as executor:
                fut_rag = executor.submit(self.create_rag_vector_store, self.resume_text)
                fut_contact = executor.submit(self.extract_contact_info, self.resume_text)
                fut_extract = executor.submit(self.extract_info_from_resume, self.resume_text)

                self.skills, self.education, self.experience = fut_extract.result()
                self.rag_vectorstore = fut_rag.result()
                self.contact_info = fut_contact.result()

            return job_id  # ONLY return job_id (no data passing)
    
    def analyze_system_new(self, role_requirements=None, custom_jd=None):
        """
        Uses internal state ONLY.
        Must be called only after preprocessing finished.
        """

        analysis = self.compare_resume_jd_new(
            skills=self.skills,
            experience=self.experience,
            education=self.education,
            role_requirements=role_requirements,
            custom_jd=custom_jd,
        )

        analysis["contact_info"] = self.contact_info

        self.analysis_result = analysis
        print("✅ contact info add")
        print("============================================================")
        print(" ")
        print("✅ everything add")
        print(" ")
        print("============================================================")
        print(" ")



        # if analysis.get("missing_skills"):
        #     analysis["detailed_weaknesses"] = self.analyze_resume_weaknesses(analysis)
        print(" ")
        print("============================================================")
        print(" ")
        print(analysis)
        print(" ")
        print("============================================================")


        return analysis, self.resume_text


    def extract_contact_info(self,text):
        # Extract email and phone using regex
        email_pattern=r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails=re.findall(email_pattern,text)
        if emails:
            self.contact_info["email"]=emails[0]

        phone_pattern=r'\b(?:\+\d{1,2}\s?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b'
        phones=re.findall(phone_pattern,text)
        if phones:
            self.contact_info["phone"]=phones[0]
        
        return self.contact_info



    def evaluate_interview(self, conversation):
        if not conversation or not isinstance(conversation, list):
            return {"error": "Invalid conversation format"}

        try:
            # Extract only the plain text from messages
            formatted_text = "\n".join(msg.get("text", "").strip() for msg in conversation if msg.get("text"))

            if len(formatted_text.strip()) < 20:
                return {"error": "Conversation too short for evaluation"}

            # Build prompt
            prompt = f"""
    You are a professional interview evaluator.

    Analyze this interview and return ONLY valid JSON.

    Format:
    {{
    "questions": [
        {{
        "question": "...",
        "answer_summary": "...",
        "score": score from 0 to 10,
        "strengths": "...",
        "improvements": "..."
        }}
    ],
    "overall_score": score from 0 to 10,
    "final_summary": "...",
    "recommendation": "Hire / Strong Hire / No Hire / Maybe"
    }}

    Interview:
    {formatted_text}
    """

            llm_response = safe_llm_invoke(self.llm, prompt)

            if not llm_response or not hasattr(llm_response, "content"):
                return {"error": "LLM returned invalid response"}

            content = llm_response.content.strip()

            # Extract JSON safely from LLM output
            match = re.search(r'\{.*\}', content, re.DOTALL)
            if not match:
                return {"error": "No JSON found in LLM response"}

            return json.loads(match.group(0))

        except Exception as e:
            return {"error": f"Evaluation failed: {str(e)}"}


    def ask_question(self, question):
        """Ask a question about the resume"""
        if not self.rag_vectorstore or not self.resume_text:
            return "Please analyze a resume first"
        
        retriever = self.rag_vectorstore.as_retriever(search_kwargs={"k": 3})
        # llm = ChatGroq(
        #     model='llama-3.1-8b-instant',
        #     api_key=GROQ_API_KEY
        # )

        
        docs = retriever.invoke(question)
        context = "\n".join([doc.page_content for doc in docs])
        
        prompt = f""" You are strict prompt follower and analyze carefully the prompt and do what only that prompt say,
        You are a helpful assistant answering questions about a resume.
        Response should be plain text only ,DONT include any special characters 
        
        Based on the following resume content only, answer the user's question accurately and concisely.
        
        Resume Content:
        {context}
        
        Question: {question}
        
        Answer:
        """
        
        response = self.llm.invoke(prompt)
        return response.content.strip() 
        

########################################################################################################################################################################################################

    def get_improved_resume(self,analysis_result):
        """Generate an improved version of the resume optimized for the job description"""
        try:
            prompt = f"""
                    YOU ARE A DETERMINISTIC RESUME FORMATTER AND COPY-EDITOR.

                    YOU ARE NOT AN IMPROVER, NOT AN OPTIMIZER, NOT A CREATOR.
                    YOU ARE ONLY ALLOWED TO COPY, LIGHTLY REPHRASE, AND REFORMAT TEXT THAT EXISTS VERBATIM IN THE RESUME.

                    IF YOU ADD EVEN ONE WORD, TOOL, METRIC, DATE, DEGREE, TITLE, SKILL, SUPERVISOR, TECHNOLOGY, OR IDEA
                    THAT DOES NOT APPEAR IN THE RESUME TEXT BELOW — THE OUTPUT IS INVALID.

                    ========================
                    SOURCE OF TRUTH — RESUME
                    ========================
                    \"\"\"{self.resume_text}\"\"\"

                    ========================
                    HARD CONSTRAINTS (ABSOLUTE)
                    ========================
                    - You MUST treat the resume text as the ONLY source of truth.
                    - You MUST NOT infer missing years, supervisors, locations, institutions, or outcomes.
                    - You MUST NOT normalize or “complete” partial information.
                    - You MUST NOT replace missing values with placeholders like [Year – Year], None, N/A, or similar.
                    - If any data is missing → DELETE THAT LINE OR ENTIRE SECTION.
                    - DO NOT mention tools, frameworks, APIs, models, metrics, cloud services, or libraries unless they are written EXACTLY in the resume.
                    - DO NOT introduce AWS, TensorFlow, supervisors, scores, rankings, accuracy, F1, percentages, or performance claims unless explicitly present.
                    - DO NOT change degree names, institute names, project titles, or event names.
                    - DO NOT summarize, extend, or generalize bullets.

                    ========================
                    CRITICAL CONTACT RULE
                    ========================
                    - CONTACT INFORMATION (email, phone, LinkedIn, GitHub) MUST APPEAR ONLY ONCE.
                    - CONTACT INFORMATION MUST APPEAR ONLY INSIDE THE CENTER BLOCK.
                    - DO NOT output contact details anywhere else.
                    - If a contact field is missing → omit ONLY that field.

                    ========================
                    SECTION CREATION GATE (CRITICAL)
                    ========================
                    A SECTION MAY BE CREATED **ONLY IF** THE RESUME TEXT CONTAINS AN EXPLICITLY LABELED SECTION
                    WITH A MATCHING OR EQUIVALENT TITLE.

                    STRICT RULES:
                    - EXPERIENCE SECTION:
                    Create ONLY if the resume explicitly contains a section labeled:
                    "Experience", "Work Experience", "Professional Experience", or "Employment".
                    PROJECT CONTENT MUST NEVER BE RECLASSIFIED AS EXPERIENCE.

                    - CERTIFICATIONS SECTION:
                    Create ONLY if the resume explicitly contains a section labeled:
                    "Certifications", "Certificates", or "Professional Certifications".
                    DO NOT treat courses, workshops, hackathons, or participation as certifications.

                    - THESIS SECTION:
                    Create ONLY if the resume explicitly contains a section labeled "Thesis" or "Dissertation".

                    - OPEN SOURCE CONTRIBUTIONS:
                    Create ONLY if explicitly labeled as "Open Source", "Open Source Contributions",
                    or similar.

                    - IF A SECTION HEADER DOES NOT EXIST IN THE RESUME → DELETE THAT ENTIRE SECTION
                    INCLUDING ITS HEADING AND RULE LINE.

                    ========================
                    ALLOWED OPERATIONS (ONLY THESE)
                    ========================
                    You may ONLY:
                    - Rephrase existing bullet points using stronger action verbs WITHOUT adding new meaning.
                    - Fix spacing, punctuation, or LaTeX escaping.
                    - Bold skills that ALREADY EXIST in the resume (\\textbf{{Skill}}).
                    - Reorder bullets within the SAME section without changing content.
                    - Convert plain text into valid LaTeX syntax.

                    ========================
                    LATEX TEMPLATE (IMMUTABLE)
                    ========================
                    You MUST use the following LaTeX structure EXACTLY.
                    DO NOT change spacing, commands, section order, or formatting.
                    ONLY replace placeholders with real extracted text OR delete the entire block.

                    \\documentclass[a4paper,12pt]{{article}}
                    \\usepackage[margin=1in]{{geometry}}
                    \\usepackage{{hyperref}}
                    \\usepackage{{enumitem}}

                    \\begin{{document}}

                    \\vspace{{-10pt}}  
                    \\begin{{center}}
                        \\Huge \\textbf{{[Candidate Name]}} \\\\ 
                        \\small
                        \\href{{mailto:[email]}}{{[email]}}
                        \\texttt{{|}} [phone]
                        \\texttt{{|}} \\href{{[linkedin_url]}}{{LinkedIn}}
                        \\texttt{{|}} \\href{{[github_url]}}{{GitHub}}
                    \\end{{center}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [SKILLS_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Technical Skills}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [EXPERIENCE_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Experience}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [PROJECTS_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Projects}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [ACHIEVEMENTS_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Achievements \\& Hackathons}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [EDUCATION_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Education}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [THESIS_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Thesis}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [CERTIFICATIONS_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Certifications}}

                    \\noindent\\rule{{\\linewidth}}{{0.4pt}}

                    [OPEN_SOURCE_SECTION]
                    \\vspace{{-10pt}}
                    \\section*{{Open Source Contributions}}

                    \\end{{document}}

                    ========================
                    FINAL OUTPUT RULE
                    ========================
                    - OUTPUT ONLY VALID, COMPILABLE LaTeX CODE.
                    - NO explanations.
                    - NO comments.
                    - NO markdown.
                    - NO placeholders.
                    - NO hallucinated content.
                    - IF A SECTION IS NOT EXPLICITLY PRESENT IN THE RESUME → DELETE IT COMPLETELY.
                    """
            try:
                # Step 1: Invoke LLM with the given prompt
                response = safe_llm_invoke(self.llm, prompt)
                
                # Step 2: Extract the LaTeX-formatted resume text
                print("✅ llm response of latex code")
                improved_resume = response.content.strip() # Remove unnecessary spaces
                # Step 3: Validate the LaTeX format
                if not improved_resume.startswith(r"\documentclass") or not improved_resume.endswith(r"\end{document}"):
                    print("❌ Generated content is not in valid LaTeX format")
                    raise ValueError("Generated content is not in valid LaTeX format.")

                # Step 5: Return the LaTeX resume content
                print("✅ generate latex code succesfully")
                print(improved_resume)
                return improved_resume

            except Exception as e:
                # Step 6: Error handling
                print(f"❌Error generating improved resume llm problem: {e}")
                return "Error generating improved resume llm problem . Please try again."
        except Exception as e:
            print(f"❌Error generating latex code: {e}")
            return "Error generating latex code. Please try again."
        

########################################################################################################################################################################################################

import streamlit as st
class Implement:
    def __init__(self):
        self.agent=ResumeAnalysisAgent()

    def analyze_resume(self,resume_file,role=None,custom_jd=None):
        """Analyze the resume with the agent"""
        return self.agent.analyze_system_new(role,custom_jd)
    
    def preprocess_resume(self,resume_file,new_job_id):
        return self.agent.preprocess_resume(resume_file,new_job_id)
    
        
    def ask_question(self,question):
        """Ask a question about the resume"""
        try:
            with st.spinner("Generating response..."):
                response=self.agent.ask_question(question)
                return response
        except Exception as e:
            return f"Error: {e}"

    def feedback_interview(self,conversation):
        """Feedback report to the Interview"""
        try:
            return self.agent.evaluate_interview(conversation)
        except Exception as e:
            return f"Error:{e}"
        
            
    def get_improved_resume(self, analysis_result):
        try:
            return self.agent.get_improved_resume(analysis_result)
        except Exception as e:
            raise RuntimeError(f"LLM resume generation failed: {e}")

